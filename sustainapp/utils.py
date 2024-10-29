from sustainapp.models import (
    Userorg,
    Organization,
    Location,
)
from django.contrib.auth.models import User
from django.db.models import Subquery
import jinja2
from django.utils import timezone
from docxtpl import DocxTemplate
from docxtpl import InlineImage
from docx.shared import Inches, Pt, RGBColor
from sustainapp.models import Report, AnalysisData2
from sustainapp.report import (
    generate_and_cache_donut_chart,
    generate_colors,
    generate_and_cache_donut_chart_source,
    generate_and_cache_donut_chart_location,
    extract_location_data,
    extract_source_data,
    word_generate_and_cache_donut_chart,
    word_generate_and_cache_donut_chart_source,
    word_generate_and_cache_donut_chart_location,
)

# Input Validations
from django.core.exceptions import ValidationError

# Email notification
from django.db.models import Q
from django.shortcuts import get_object_or_404
from django.contrib.staticfiles import finders
import pycountry
import datetime
import json
from docx import Document
from docx.shared import Inches
from django.core.files import File
from azure.storage.blob import BlobServiceClient
from django.core.files.storage import default_storage
import os
import io
from django.http import StreamingHttpResponse, HttpResponse
from html.parser import HTMLParser
from htmldocx import HtmlToDocx
import logging
from rest_framework.views import exception_handler
from rest_framework.response import Response
from rest_framework import status
from django.core.exceptions import PermissionDenied
from django.conf import settings
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from common.utils.value_types import get_decimal

logger = logging.getLogger(__name__)


def client_request_data(data, client_obj):

    data["client_id"] = client_obj.id
    data["client"] = client_obj

    return data


def client_request_data_modelviewset(data, client_obj):

    data["client_id"] = client_obj.id
    data["client"] = client_obj.id

    return data


def check_same_client(client, instance):
    request_client_id = getattr(client, "id", None)
    instance_client_id = instance.client_id if instance and instance.client else None
    if request_client_id != instance_client_id:
        logging.error(
            f"Client ID mismatch: request client ID {request_client_id},instance client ID {instance_client_id}"
        )
        raise PermissionDenied("You do not have permission to access this resource")


def custom_exception_handler(exc, context):
    response = exception_handler(exc, context)

    # Now add custom handling
    if isinstance(exc, PermissionDenied):
        return Response({"error": str(exc)}, status=status.HTTP_403_FORBIDDEN)

    # If it's not handled above, use the default handler
    return response


def validation_get_emission(data):
    required_fields = ["location", "year", "month", "client_id"]

    missing_fields = [field for field in required_fields if field not in data]

    if missing_fields:
        raise ValidationError(f"Missing fields: {', '.join(missing_fields)}")


def validation_org_input(data):
    required_fields = [
        "name",
        "owner",
        "countryoperation",
        "employeecount",
        "website",
        "active",
        "type_corporate_entity",
        "location_of_headquarters",
        "phone",
        "mobile",
        "fax",
        "sector",
        "revenue",
        "subindustry",
        "address",
        "state",
        "city",
        "date_format",
        "currency",
        "timezone",
        "language",
        "from_date",
        "to_date",
        "framework",
        "client",
    ]
    missing_fields = [field for field in required_fields if field not in data]

    if missing_fields:
        raise ValidationError(f"Missing fields: {', '.join(missing_fields)}")


def validation_location_input(data):
    required_fields = [
        "corporateentity_id",
        "corporateentity",
        "name",
        "area",
        "city",
        "country",
        "currency",
        "dateformat",
        "employeecount",
        "fax",
        "from_date",
        "language",
        "latitude",
        "location_type",
        "longitude",
        "mobile",
        "name",
        "phone",
        "revenue",
        "sector",
        "state",
        "streetaddress",
        "sub_industry",
        "timezone",
        "to_date",
        "type_of_business_activities",
        "type_of_product",
        "type_of_services",
        "typelocation",
        "website",
        "zipcode",
        "longitude",
        "latitude",
    ]
    missing_fields = [field for field in required_fields if field not in data]
    if missing_fields:
        raise ValidationError(f"Missing fields: {', '.join(missing_fields)}")


def validation_corporate_input(data):
    required_fields = [
        "name",
        "corporatetype",
        "ownershipnature",
        "location_headquarters",
        "phone",
        "mobile",
        "website",
        "fax",
        "employeecount",
        "revenue",
        "sector",
        "subindustry",
        "address",
        "Country",
        "state",
        "city",
        "zipcode",
        "date_format",
        "currency",
        "timezone",
        "language",
        "from_date",
        "to_date",
        "legalform",
        "ownership",
        "group",
        "location_of_headquarters",
        "amount",
        "framework",
        "organization_id",
        "type_of_business_activities",
        "type_of_product",
    ]
    # "type_of_services",
    # "type_of_business_relationship",
    # "framework","organization_id"]
    missing_fields = [field for field in required_fields if field not in data]
    if missing_fields:
        raise ValidationError(f"Missing fields: {', '.join(missing_fields)}")


def org_user(username, *args, **kwargs):
    user_org = list(
        Userorg.objects.filter(
            user__in=Subquery(User.objects.filter(username=username).values("id"))
        ).values("organization")
    )[0]["organization"]
    organization = list(Organization.objects.filter(id=user_org).values("name"))[0][
        "name"
    ]
    return organization


def location_organization(organization, *args, **kwargs):
    return [
        individual_ids["id"]
        for individual_ids in list(
            Location.objects.filter(
                corporateentity__organization__name=organization
            ).values("id")
        )
    ]


def org_from_location(Location):
    organization = (
        Organization.objects.filter(corporatenetityorg__location__name=Location)
        .values_list("name", flat=True)
        .first()
    )
    return organization


def handle_none(value):
    return 0 if value is None else value


def creating_task_status(queryset):
    final_data = {"upcoming": "", "overdue": "", "completed": ""}
    today = timezone.now().date()
    final_data["completed"] = queryset.filter(completed=True)
    final_data["upcoming"] = queryset.filter(Q(completed=False, deadline__gte=today))
    final_data["overdue"] = queryset.filter(Q(completed=False, deadline__lt=today))

    return final_data


def html_to_plain_text(html):
    from urllib.request import urlopen


class HtmlElementParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.elements = []
        self.current_element = None
        self.current_row = None
        self.current_cell = None

    def handle_starttag(self, tag, attrs):
        if tag == "p":
            self.current_element = {"type": "paragraph", "content": ""}
        elif tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            self.current_element = {
                "type": "heading",
                "level": int(tag[1]),
                "content": "",
            }
        elif tag == "tr":
            self.current_row = []
            self.current_cell = ""
        elif tag == "td":
            self.current_cell = ""

    def handle_data(self, data):
        if self.current_element:
            self.current_element["content"] += data
        elif self.current_cell is not None:
            self.current_cell += data.strip()

    def handle_endtag(self, tag):
        if tag == "p":
            if self.current_element["content"]:
                self.elements.append(self.current_element)
            self.current_element = None
        elif tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            self.elements.append(self.current_element)
            self.current_element = None
        elif tag == "tr":
            if self.current_row:
                self.elements.append({"type": "table_row", "content": self.current_row})
            self.current_row = None
        elif tag == "td":
            if self.current_cell:
                self.current_row.append(self.current_cell)
            self.current_cell = None

    def handle_startendtag(self, tag, attrs):
        if tag == "br":
            if self.current_element:
                self.current_element["content"] += "\n"
            elif self.current_cell is not None:
                self.current_cell += "\n"


def parse_html_elements(html_content):
    parser = HtmlElementParser()
    parser.feed(html_content)
    return parser.elements


def generate_report_data(pk, request):
    report = Report.objects.get(id=pk)
    country_code = report.organization.countryoperation

    if not country_code:
        country_name = "Unknown"
    else:
        try:
            # Attempt to get the country name from pycountry
            country_name = pycountry.countries.get(alpha_2=country_code).name
        except AttributeError:
            # Handle the case where the country_code is not found in pycountry
            country_name = "Unknown"

    image_path = finders.find("images/ghg-methodology-flowchart.png")

    data_entry = get_object_or_404(AnalysisData2, report_id=pk)
    try:
        data_dict = json.loads(data_entry.data.replace('\\"', '"'))
    except json.JSONDecodeError as e:
        print(f"JSON Decode Error: {e}")

    # data_dict = model_to_dict(data_entry)
    organized_data_list = []
    total_co2e_combined = 0
    total_contribution_combined = 0
    combined_scopes = {}

    # Iterate over each corporate in data_dict
    for corporate_name, corporate_data in data_dict.items():
        # Extract data for each corporate
        scopes = corporate_data.get("scope", [])
        locations = corporate_data.get("location", [])
        sources = corporate_data.get("source", [])

        # Calculate total combined CO2e and total contribution for each corporate

        total_co2e_corporate = sum(float(scope["total_co2e"]) for scope in scopes)
        total_co2e_combined += total_co2e_corporate

        # Calculate total contribution for all corporates
        total_contribution_corporate = sum(
            float(scope["contribution_scope"]) for scope in scopes
        )
        total_contribution_combined += total_contribution_corporate

        # Calculate combined percentage for each scope within the corporate
        for scope in scopes:
            scope_name = scope["scope_name"]

            if scope_name not in combined_scopes:
                combined_scopes[scope_name] = {
                    "scope_name": scope_name,
                    "total_co2e": float(scope["total_co2e"]),
                    "contribution_scope": float(scope["contribution_scope"]),
                    "co2e_unit": scope["co2e_unit"],
                }
            else:
                combined_scopes[scope_name]["total_co2e"] += float(scope["total_co2e"])
                combined_scopes[scope_name]["combined_percentage"] = (
                    (
                        float(combined_scopes[scope_name]["total_co2e"])
                        / total_co2e_combined
                    )
                    * 100
                    if total_co2e_combined != 0
                    else 0
                )

        # After updating total_co2e for all scopes in the current corporate, calculate combined_percentage
        for scope_name, scope_data in combined_scopes.items():
            scope_data["combined_percentage"] = (
                (scope_data["total_co2e"] / total_co2e_combined) * 100
                if total_co2e_combined != 0
                else 0
            )
            # Convert the combined scopes back to a list if necessary
            combined_scopes_list = list(combined_scopes.values())

        # Organize the data
        organized_data = {
            "corporate_name": corporate_name,
            "scopes": scopes,
            "locations": locations,
            "sources": sources,
        }

        organized_data_list.append(organized_data)
        highest_contribution_value = 0
        highest_source_name = None
        for item in organized_data_list:
            highest_contribution_source_for_corporate = max(
                item["sources"],
                key=lambda x: float(x["contribution_source"]),
            )
            # Compare the numeric value instead of the dictionary
            if (
                float(highest_contribution_source_for_corporate["contribution_source"])
                > highest_contribution_value
            ):
                highest_contribution_value = float(
                    highest_contribution_source_for_corporate["contribution_source"]
                )
                highest_source_name = highest_contribution_source_for_corporate[
                    "source_name"
                ]

    # Extract total_co2e and scope_name from combined_scopes
    extracted_data = [
        {"scope_name": scope["scope_name"], "total_co2e": scope["total_co2e"]}
        for scope in combined_scopes.values()
    ]
    # Call the function to generate donut chart HTML
    donut_chart_html = generate_and_cache_donut_chart(extracted_data)

    # report_list_view, extract source data
    source_data = extract_source_data(organized_data_list)

    # Use the extracted source_data to generate the dynamic chart
    source_donut_chart_html = generate_and_cache_donut_chart_source(source_data)
    source_donut_chart_image, legend_scouce_donut_chart_image = (
        word_generate_and_cache_donut_chart_source(source_data)
    )
    # report_list_view, extract source data
    location_data = extract_location_data(organized_data_list)

    # Use the extracted source_data to generate the dynamic chart
    location_donut_chart_html = generate_and_cache_donut_chart_location(location_data)
    location_donut_chart_image, legend_location_donut_chart_image = (
        word_generate_and_cache_donut_chart_location(location_data)
    )

    donut_chart_image, legend_donut_chart_image = word_generate_and_cache_donut_chart(
        extracted_data
    )
    new_parser = HtmlToDocx()

    about_the_organization = (
        new_parser.parse_html_string(report.about_the_organization)
        if report.about_the_organization is not None
        else None
    )

    roles_and_responsibilities = (
        new_parser.parse_html_string(report.roles_and_responsibilities)
        if report.roles_and_responsibilities is not None
        else None
    )

    excluded_sources = (
        new_parser.parse_html_string(report.excluded_sources)
        if report.excluded_sources is not None
        else None
    )
    organization_name = report.organization.name
    if report.report_by == "Organization":
        organization_name = organization_name
    else:
        organization_name = report.corporate.name

    context = {
        "object_list": report,
        "org_logo": report.org_logo,
        "report_by": report.report_by,
        "report_type": report.report_type,
        "organization_name": organization_name,
        "country": country_name,
        "about_the_organization": about_the_organization,
        "bool_about_the_organization": bool(about_the_organization),
        "roles_and_responsibilities": roles_and_responsibilities,
        "bool_roles_and_responsibilities": bool(roles_and_responsibilities),
        "organizational_boundries": report.organizational_boundries,
        "excluded_sources": excluded_sources,
        "bool_excluded_sources": bool(excluded_sources),
        "designation_of_organizational_admin": report.designation_of_organizational_admin,
        "reporting_period_name": report.reporting_period_name,
        "start_date": report.start_date,
        "end_date": report.end_date,
        "from_year": report.from_year,
        "to_year": report.to_year,
        "data_source": report.data_source,
        "calender_year": report.calender_year,
        "total_co2e_combined": round(total_co2e_combined, 2),
        "pk": pk,
        "data": organized_data_list,
        "combined_scopes": combined_scopes_list,
        "image_path": image_path,
        "donut_chart_html": donut_chart_html,
        "legend_donut_chart_image": legend_donut_chart_image,
        "source_donut_chart_html": source_donut_chart_html,
        "location_donut_chart_html": location_donut_chart_html,
        "scope_donut_chart_image": donut_chart_image,
        "source_donut_chart_image": source_donut_chart_image,
        "legend_scouce_donut_chart_image": legend_scouce_donut_chart_image,
        "location_donut_chart_image": location_donut_chart_image,
        "legend_location_donut_chart_image": legend_location_donut_chart_image,
        "highest_source_name": highest_source_name,
    }

    return context


def adjust_paragraphs(doc_stream):
    try:
        doc_stream.seek(0)
        doc = Document(doc_stream)
        all_paragraphs = list(doc.paragraphs)

        for index, paragraph in enumerate(all_paragraphs):
            text = paragraph.text.strip()
            if not text:
                # Remove empty paragraphs completely
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None
            else:
                # Remove spacing for first and last paragraphs
                if index == 0:
                    paragraph.paragraph_format.space_before = Pt(0)
                if index == len(all_paragraphs) - 1:
                    paragraph.paragraph_format.space_after = Pt(0)
                # Standard space after for other paragraphs
                else:
                    paragraph.paragraph_format.space_after = Pt(12)

        modified_stream = io.BytesIO()
        doc.save(modified_stream)
        modified_stream.seek(0)
        return modified_stream
    except Exception as e:
        logger.error(f"Error adjusting paragraphs: {e}")
        return doc_stream


def word_docx_report(pk: int):
    blob_name = os.path.join(
        settings.MEDIA_ROOT, "files", "report", "report_demo_v2.docx"
    )
    blob_object = default_storage.open(blob_name, "rb")
    blob_stream = io.BytesIO(blob_object.read())
    tpl = DocxTemplate(blob_stream)
    context = generate_report_data(pk, None)
    about_the_organization_stream = io.BytesIO()
    roles_and_responsibilities_stream = io.BytesIO()
    excluded_sources_stream = io.BytesIO()
    if context["about_the_organization"]:
        context["about_the_organization"].save(about_the_organization_stream)
        adjusted_about_the_organization_stream = adjust_paragraphs(
            about_the_organization_stream
        )
        context["about_the_organization"] = tpl.new_subdoc(
            adjusted_about_the_organization_stream
        )
    if context["excluded_sources"]:
        context["excluded_sources"].save(excluded_sources_stream)
        adjusted_excluded_sources_stream = adjust_paragraphs(excluded_sources_stream)
        context["excluded_sources"] = tpl.new_subdoc(adjusted_excluded_sources_stream)

    if context["roles_and_responsibilities"]:
        context["roles_and_responsibilities"].save(roles_and_responsibilities_stream)
        adjusted_roles_and_responsibilities_stream = adjust_paragraphs(
            roles_and_responsibilities_stream
        )
        context["roles_and_responsibilities"] = tpl.new_subdoc(
            adjusted_roles_and_responsibilities_stream
        )
    if not bool(context["org_logo"]):
        context["org_logo"] = default_storage.open("sustainext.jpeg")
    context["org_logo"] = io.BytesIO(context["org_logo"].read())
    inline_image = InlineImage(
        tpl=tpl,
        image_descriptor=context["org_logo"],
        width=Inches(2.5),
    )
    context["org_logo"] = inline_image
    jinja_env = jinja2.Environment(autoescape=True)

    tpl.render(context, jinja_env)
    output_bytes_object = io.BytesIO()
    tpl.save(output_bytes_object)
    document_object = Document(output_bytes_object)
    # add_page_break_if_needed(document_object)

    # Emissions by Scope
    scope_heading = document_object.add_heading(level=5).add_run("Emissions by Scope")
    scope_heading.font.size = Pt(14)
    scope_heading.bold = True
    scope_heading.font.color.rgb = RGBColor(0, 0, 0)

    scope_table = document_object.add_table(rows=2, cols=1)
    scope_table.style = "Table Grid"
    for cell in scope_table.columns[0].cells:
        set_cell_border(cell, border_color="#FFFFFF", border_width=0)
    scope_top_cell = scope_table.cell(0, 0)
    scope_bottom_cell = scope_table.cell(1, 0)
    scope_top_paragraph = scope_top_cell.paragraphs[0]
    scope_bottom_paragraph = scope_bottom_cell.paragraphs[0]
    scope_top_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope_bottom_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    scope_chart_run = scope_top_paragraph
    scope_chart_run.add_run().add_picture(
        context["scope_donut_chart_image"], width=Inches(3.5)
    )
    scope_legend_run = scope_bottom_paragraph
    scope_legend_run.add_run().add_picture(context["legend_donut_chart_image"])

    # Emissions by Source
    # Add headings and images using a table for "Emissions by Source"
    source_heading = document_object.add_heading(level=5).add_run("Emissions by Source")
    source_heading.font.size = Pt(14)
    source_heading.bold = True
    source_heading.font.color.rgb = RGBColor(0, 0, 0)

    # Insert images using a table for "Emissions by Source"
    table = document_object.add_table(rows=2, cols=1)
    set_table_keep_together(table)
    table.style = "Table Grid"
    for cell in table.columns[0].cells:
        set_cell_border(cell, border_color="#FFFFFF", border_width=0)
    left_cell = table.cell(0, 0)  # For the chart image
    right_cell = table.cell(1, 0)  # For the legend image

    left_paragraph = left_cell.paragraphs[0]
    right_paragraph = right_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add the chart image
    chart_run = left_paragraph.add_run()
    chart_run.add_picture(context["source_donut_chart_image"], width=Inches(3.5))

    # Add the legend image
    legend_run = right_paragraph.add_run()
    legend_run.add_picture(context["legend_scouce_donut_chart_image"])

    # Emissions by Location
    location_heading = document_object.add_heading(level=5).add_run(
        "Emissions by Location"
    )
    location_heading.font.size = Pt(14)
    location_heading.bold = True
    location_heading.font.color.rgb = RGBColor(0, 0, 0)

    location_table = document_object.add_table(rows=2, cols=1)
    set_table_keep_together(location_table)
    location_table.style = "Table Grid"
    for cell in location_table.columns[0].cells:
        set_cell_border(cell, border_color="#FFFFFF", border_width=0)
    location_top_cell = location_table.cell(0, 0)
    location_bottom_cell = location_table.cell(1, 0)
    location_top_paragraph = location_top_cell.paragraphs[0]
    location_bottom_paragraph = location_bottom_cell.paragraphs[0]
    location_top_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    location_bottom_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    location_chart_run = location_top_paragraph
    location_chart_run.add_run().add_picture(
        context["location_donut_chart_image"], width=Inches(3.5)
    )
    location_legend_run = location_bottom_paragraph
    location_legend_run.add_run().add_picture(
        context["legend_location_donut_chart_image"]
    )

    output_model_bytes_object = io.BytesIO()

    file_name = f"{context['object_list'].name}.docx"
    document_object.save(output_model_bytes_object)
    output_model_bytes_object.seek(0)
    document_content = (
        output_model_bytes_object.read()
    )  # <-- To convert to bytes and send as single response
    response = HttpResponse(
        content=document_content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response["Content-Disposition"] = f"attachment;filename={file_name}"
    response["Content-Encoding"] = "UTF-8"

    return response


def set_cell_border(cell, border_color="auto", border_width=4, border_space=0):
    """Set border style for a cell in a Word table."""
    tc_pr = cell._element.get_or_add_tcPr()  # Get or add the tcPr element

    # Specify the borders you want to modify
    borders = {"top": "single", "left": "single", "bottom": "single", "right": "single"}

    for key, value in borders.items():
        tag = f"w:{key}"
        border_el = OxmlElement(tag)  # Create element for each border side
        border_el.set(qn("w:val"), value)  # Type of border
        border_el.set(qn("w:sz"), str(border_width))  # Size of border
        border_el.set(qn("w:color"), border_color)  # Border color
        border_el.set(
            qn("w:space"), str(border_space)
        )  # Space between borders and text

        tc_pr.append(border_el)


def set_table_keep_together(table):
    """Set table properties to keep the table together on one page."""
    tbl = table._tbl  # Get the underlying <w:tbl> XML element
    tblPr = tbl.find(qn("w:tblPr"))  # Try to find the <w:tblPr> element
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")  # Create a new <w:tblPr> element if not found
        tbl.insert(0, tblPr)  # Insert it at the beginning of the table element

    keepNext = OxmlElement("w:keepNext")  # Create <w:keepNext/> element
    keepNext.set(qn("w:val"), "1")  # Set attribute w:val="1"
    tblPr.append(keepNext)  # Append to the <w:tblPr> element

    keepLines = OxmlElement("w:keepLines")  # Create <w:keepLines/> element
    keepLines.set(qn("w:val"), "1")  # Set attribute w:val="1"
    tblPr.append(keepLines)  # Append to the <w:tblPr> element


def add_page_break_if_needed(document):
    """Add a page break if the remaining space is not enough."""
    # This function is very basic, as python-docx does not provide a way to measure remaining space.
    # You might need a heuristic or manual trigger to decide when to add a page break.
    document.add_page_break()


def cell_transparent(table_name):
    for cell in table_name.columns[0].cells:
        set_cell_border(cell, border_color="#FFFFFF", border_width=0)


def get_ratio_of_annual_total_compensation_ratio_of_percentage_increase_in_annual_total_compensation(
    raw_response, slugs: dict
):
    annual_raw_response = raw_response.only("data").filter(path__slug=slugs[0])
    contextual_raw_response = raw_response.only("data").filter(path__slug=slugs[1])

    local_annual_response = []
    for raw_response in annual_raw_response:
        local_annual_response.extend(raw_response.data)
    local_contextual_response = []
    for raw_response in contextual_raw_response:
        local_contextual_response.extend(raw_response.data)
    local_response_data = []
    for annual, contextual in zip(local_annual_response, local_contextual_response):
        local_response_data.append(
            {
                "ratio_of_annual_total_compensation": get_decimal(
                    get_decimal(annual["Q1"]) / get_decimal(annual["Q2"])
                ),
                "ratio_of_percentage_increase_in_annual_total_compensation": get_decimal(
                    get_decimal(contextual["Q1"]) / get_decimal(contextual["Q2"])
                ),
            }
        )
    return local_response_data


def validation_for_esg_report(report: Report):
    paths_and_disclosures = [
        {
            "general": [
                {
                    "GRI Reporting Info": {
                        "Org Details": [
                            (
                                "Organizational Details",
                                "gri-general-org_details_2-1a-1b-1c-1d",
                            )
                        ],
                        "Entities": [
                            (
                                "List of entities",
                                "gri-general-entities-list_of_entities-2-2-a",
                            ),
                            (
                                "Audited, consolidated financial statements",
                                "gri-general-entities-audited-2-2-b",
                            ),
                            (
                                "Multiple entities",
                                "gri-general-entities-multiple-2-2-c",
                            ),
                        ],
                        "Report Details": (),
                    }
                },
            ]
        }
    ]


def create_users():
    from authentication.models import CustomUser

    def bulk_create_users(email_list):
        users_to_create = []
        for email in email_list:
            user = CustomUser(
                username=email,
                email=email,
                work_email=email,
                is_active=True,
                is_staff=True,
                is_superuser=True,
            )
            users_to_create.append(user)

        return CustomUser.objects.bulk_create(users_to_create)

    emails = [
        "ajay.korpal@sustainext.ai",
        "onkar.bhave@sustainext.ai",
        "shubham.kanungo@sustainext.ai",
        "anush.hp@sustainext.ai",
        "sahana.kv@sustainext.ai",
        "utsav.pipersaniya@sustainext.ai",
        "himanshu.banswal@sustainext.ai",
        "sakthivel.murugan@sustainext.ai",
        "yashwanth.vanimina@sustainext.ai",
        "mahinder.singh@sustainext.ai",
    ]
    created_users = bulk_create_users(emails)
    from allauth.account.models import EmailAddress, EmailConfirmation
    for user in CustomUser.objects.all():
        user.set_password("sustainext@1234")
        EmailAddress.objects.create(
            user=user, email=user.email, verified=True, primary=True
        ).save()

import json
import logging
from django.core.exceptions import ValidationError
from django.db import transaction


logger = logging.getLogger(__name__)


def load_successful_raw_responses_from_json(file_path):
    from sustainapp.models import Corporateentity, Organization, Location
    """
    Load RawResponse objects from a JSON fixture, skipping entries with signal errors.

    Args:
        file_path (str): Path to the JSON fixture file.

    Returns:
        list: Successfully loaded RawResponse instances.
    """
    successful_responses = []

    # Read JSON data from fixture file
    with open(file_path, "r") as f:
        data = json.load(f)

    # Iterate over each entry in the JSON data
    for raw_data in data:
        try:
            with transaction.atomic():  # Atomic block to prevent partial saves
                # Remove "model" and "pk" fields used in Django's fixture format
                fields = raw_data["fields"]
                fields["path"] = Path.objects.get(id=fields["path"])
                fields["user"] = CustomUser.objects.get(id=fields["user"])
                fields["client"] = Client.objects.get(id=fields["client"])
                if fields["organization"]:
                    fields["organization"] = Organization.objects.get(
                        id=fields["organization"]
                    )
                if fields["corporate"]:
                    fields["corporate"] = Corporateentity.objects.get(
                        id=fields["corporate"]
                    )
                if fields["locale"]:
                    fields["locale"] = Location.objects.get(id=fields["locale"])
                fields["month"] = fields["month"] if fields["month"] else None
                fields["year"] = fields["year"] if fields["year"] else None
                response = RawResponse(**fields)
                response.full_clean()  # Validate all fields
                response.save()
                successful_responses.append(response)

        except ValidationError as e:
            # Log validation errors or handle as needed
            print(f"Validation error loading RawResponse data: {e}")

        except Exception as e:
            # Catch other exceptions (likely from signals) and log them
            print(f"Error loading RawResponse object: {e}")
            continue
